from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT = r"C:\Users\julia\Desktop\Shooter1persona\documentacion"

authors = "Andres Paolo Casta\u00f1o Velez\nCristian Giovanni Castrillon Arias\nJulio Cesar Caicedo Eraso"
faculty = "Facultad de inteligencia artificial e ingenier\u00eda\nUniversidad de Caldas\n2025"

def save(doc, name):
    doc.save(os.path.join(OUT, name))
    print(f"Created {name}")

def add_title_page(doc, title, subtitle=""):
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Shooter1persona")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(22)
        r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    for _ in range(4):
        doc.add_paragraph("")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(authors).font.size = Pt(13)
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(faculty).font.size = Pt(12)
    doc.add_page_break()

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x55)

def para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

# ================================================================
# 1. INSTALACION
# ================================================================
doc = Document()
add_title_page(doc, "Shooter1persona", "Manual de Instalaci\u00f3n")

h1(doc, "Descripci\u00f3n")
para(doc, "Shooter1persona es un videojuego de disparos en primera persona (FPS) desarrollado en Unity 6000.3.7f1 con URP 17.3.0. El jugador debe sobrevivir a rondas de enemigos en un entorno 3D, utilizando un arma con sistema de munici\u00f3n y recarga, barra de vida con Slider, y monedas de vida para recuperar salud. Cada ronda incrementa la cantidad de enemigos, aumentando la dificultad progresivamente.")

h1(doc, "Requisitos del Sistema")
h2(doc, "M\u00ednimos")
for req in [
    "Sistema Operativo: Windows 10 (64-bit)",
    "Procesador: Intel Core i3-2100 o AMD FX-6300",
    "Memoria RAM: 8 GB",
    "Tarjeta Gr\u00e1fica: NVIDIA GeForce GTX 660 o AMD Radeon HD 7870",
    "DirectX: Versi\u00f3n 11",
    "Almacenamiento: 1 GB de espacio disponible"
]:
    doc.add_paragraph(req, style='List Bullet')

h2(doc, "Recomendados")
for req in [
    "Sistema Operativo: Windows 11 (64-bit)",
    "Procesador: Intel Core i5-8400 o AMD Ryzen 5 2600",
    "Memoria RAM: 16 GB",
    "Tarjeta Gr\u00e1fica: NVIDIA GeForce GTX 1060 o AMD Radeon RX 580",
    "DirectX: Versi\u00f3n 12",
    "Almacenamiento: 2 GB de espacio disponible"
]:
    doc.add_paragraph(req, style='List Bullet')

h1(doc, "Instalaci\u00f3n")
para(doc, "Al ser un proyecto desarrollado en Unity, el proceso de instalaci\u00f3n es sencillo:")

h2(doc, "Desde el ejecutable")
pasos = [
    "Ubicar la carpeta 'Build' o 'Ejecutable' dentro del proyecto.",
    "Ejecutar el archivo 'Shooter1persona.exe'.",
    "Si aparece una advertencia de seguridad de Windows, seleccionar 'Ejecutar de todas formas'.",
    "Esperar a que cargue la escena de men\u00fa principal."
]
for i, paso in enumerate(pasos, 1):
    doc.add_paragraph(f"Paso {i}: {paso}", style='List Number')

h2(doc, "Desde el Editor de Unity")
pasos_editor = [
    "Abrir el proyecto en Unity 6000.3.7f1.",
    "En la ventana 'Build Settings', seleccionar la plataforma 'PC, Mac & Linux Standalone'.",
    "Asegurarse de que las escenas 'MenuInicio' y 'SampleScene' est\u00e9n a\u00f1adidas en 'Scenes In Build'.",
    "Hacer clic en 'Build And Run'.",
    "Seleccionar la carpeta de destino y esperar a que se compile el ejecutable."
]
for i, paso in enumerate(pasos_editor, 1):
    doc.add_paragraph(f"Paso {i}: {paso}", style='List Number')

save(doc, "Instalacion.docx")

# ================================================================
# 2. TECNICO
# ================================================================
doc = Document()
add_title_page(doc, "Shooter1persona", "Manual T\u00e9cnico")

h1(doc, "Metodolog\u00eda de Desarrollo")
para(doc, "El proyecto corresponde a un desarrollo tecnol\u00f3gico fundamentado en t\u00e9cnicas de ingenier\u00eda de software aplicadas a un motor de desarrollo de videojuegos.")
para(doc, "Se utiliz\u00f3 la metodolog\u00eda \u00e1gil ICONIX, que se enfoca en pasar de los casos de uso al c\u00f3digo a trav\u00e9s de un sistema de an\u00e1lisis y dise\u00f1o, utilizando un subconjunto esencial de UML con solo cuatro tipos de diagramas, optimizando tiempo y recursos.")

h1(doc, "Herramientas y Tecnolog\u00eda")
h2(doc, "Unity 6000.3.7f1")
para(doc, "Motor de desarrollo para crear contenido 3D interactivo. Se utiliz\u00f3 con el pipeline Universal Render Pipeline (URP) 17.3.0 para optimizar el rendimiento gr\u00e1fico.")
h2(doc, "C#")
para(doc, "Lenguaje de programaci\u00f3n orientado a objetos desarrollado por Microsoft, utilizado para la l\u00f3gica del juego, incluyendo movimiento del jugador, sistema de disparo, inteligencia artificial de enemigos y gesti\u00f3n de la interfaz de usuario.")
h2(doc, "Bibliotecas y Paquetes")
for pkg in [
    "com.unity.ai.navigation 2.0.9 - Navegaci\u00f3n de enemigos con NavMesh",
    "com.unity.inputsystem 1.18.0 - Sistema de entrada moderno",
    "com.unity.render-pipelines.universal 17.3.0 - Pipeline de renderizado URP",
    "com.unity.test-framework 1.6.0 - Framework de pruebas unitarias",
    "com.unity.textmeshpro - Texto renderizado para HUD",
    "com.unity.ugui 2.0.0 - Interfaz de usuario"
]:
    doc.add_paragraph(pkg, style='List Bullet')

h1(doc, "Arquitectura del Sistema")
para(doc, "El proyecto est\u00e1 organizado en una escena de men\u00fa principal (MenuInicio.unity) y una escena de juego (SampleScene.unity), ambas gestionadas mediante SceneManager.")

h2(doc, "Estructura de Scripts")
scripts_info = [
    ("FPSController.cs", "Control del jugador: movimiento WASD, sprint con Shift, mirada con el rat\u00f3n mediante CharacterController."),
    ("Arma.cs", "Sistema de disparo por raycast, efecto de muzzle flash, estela con LineRenderer, detecci\u00f3n de enemigo en la mira, sistema de munici\u00f3n con cargador (30 balas), reserva total (120 balas) y recarga con tecla R."),
    ("Enemigo.cs", "Inteligencia artificial con NavMeshAgent, persecuci\u00f3n del jugador, ataque cuerpo a cuerpo, posibilidad de mutaci\u00f3n (30%) a variante peque\u00f1a y r\u00e1pida, muerte con drops."),
    ("JugadorSalud.cs", "Gesti\u00f3n de salud del jugador (150 HP), da\u00f1o, curaci\u00f3n, barra de vida con Slider (en lugar de Image.FillMethod), pantalla de muerte y reinicio de partida."),
    ("SpawnerEnemigos.cs", "Sistema de rondas progresivas: cada ronda aumenta la cantidad de enemigos (base + incremento por ronda), pausa entre rondas con anuncio visual, contador de bajas con TextMeshPro."),
    ("GestorPausa.cs", "Pausa con Escape, men\u00fa completo con opciones Reanudar, Reiniciar (recarga la escena) y Salir (Application.Quit)."),
    ("BalanceoArma.cs", "Animaci\u00f3n de balanceo del arma al caminar y correr."),
    ("SeguirJugadorMinimapa.cs", "C\u00e1mara del minimapa que sigue al jugador bloqueada en el eje Y."),
    ("MonedaVida.cs", "Objeto de curaci\u00f3n que se recoge al colisionar, recupera 20 HP."),
    ("MenuPrincipal.cs", "Men\u00fa principal con opciones Iniciar Juego y Salir.")
]
for name, desc in scripts_info:
    doc.add_paragraph(f"{name}: {desc}", style='List Bullet')

h1(doc, "Requerimientos del Sistema")
h2(doc, "Requerimientos Funcionales")
table = doc.add_table(rows=9, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['C\u00f3digo', 'Requerimiento']
data = [
    ('RFS01', 'El jugador debe poder desplazarse por el entorno 3D con WASD y correr con Shift.'),
    ('RFS02', 'El jugador debe poder apuntar y disparar con el clic izquierdo del rat\u00f3n.'),
    ('RFS03', 'Los enemigos deben perseguir al jugador mediante NavMesh.'),
    ('RFS04', 'El sistema debe gestionar la salud del jugador y mostrar una barra de vida con Slider.'),
    ('RFS05', 'El sistema debe generar rondas progresivas de enemigos. Cada ronda aumenta la cantidad de enemigos respecto a la anterior.'),
    ('RFS06', 'El juego debe contar con un men\u00fa de pausa (Reanudar, Reiniciar, Salir).'),
    ('RFS07', 'El jugador debe poder recoger monedas de vida para recuperar salud.'),
    ('RFS08', 'El arma debe tener un sistema de munici\u00f3n con cargador y recarga manual (tecla R).'),
]
for j, h in enumerate(headers):
    table.rows[0].cells[j].text = h
for i, (code, req) in enumerate(data):
    table.rows[i+1].cells[0].text = code
    table.rows[i+1].cells[1].text = req

doc.add_paragraph()

h2(doc, "Requerimientos No Funcionales")
table2 = doc.add_table(rows=6, cols=2)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(headers):
    table2.rows[0].cells[j].text = h
data_nf = [
    ('RNF01', 'El sistema debe funcionar en tiempo real con una frecuencia m\u00ednima de 30 FPS.'),
    ('RNF02', 'La interfaz debe actualizarse en tiempo real (barra de vida, contador de bajas, munici\u00f3n, ronda actual).'),
    ('RNF03', 'El cursor debe bloquearse durante el juego y liberarse en pausa o muerte.'),
    ('RNF04', 'El sistema debe cargar la escena de juego desde el men\u00fa principal.'),
    ('RNF05', 'El sistema de rondas debe mostrar un anuncio visual al inicio de cada ronda.'),
]
for i, (code, req) in enumerate(data_nf):
    table2.rows[i+1].cells[0].text = code
    table2.rows[i+1].cells[1].text = req

doc.add_paragraph()

h1(doc, "Diagramas de Casos de Uso")

h2(doc, "Caso de Uso: Men\u00fa Principal")
para(doc, "Actor: Jugador")
para(doc, "Descripci\u00f3n: El jugador accede al men\u00fa principal y elige entre iniciar el juego o salir.")
para(doc, "Precondici\u00f3n: La aplicaci\u00f3n se ha iniciado correctamente.")
para(doc, "Secuencia normal:")
for i, s in enumerate([
    "El jugador visualiza las opciones del men\u00fa.",
    "Selecciona 'Jugar' para iniciar la partida.",
    "Selecciona 'Salir' para cerrar la aplicaci\u00f3n."
], 1):
    doc.add_paragraph(f"{i}. {s}", style='List Number')

h2(doc, "Caso de Uso: Juego")
para(doc, "Actor: Jugador")
para(doc, "Descripci\u00f3n: El jugador se desplaza por el escenario, dispara a los enemigos, gestiona su salud y su munici\u00f3n.")
para(doc, "Precondici\u00f3n: La escena SampleScene se ha cargado correctamente.")
para(doc, "Secuencia normal:")
for i, s in enumerate([
    "El jugador se desplaza con WASD y mira con el rat\u00f3n.",
    "Comienza la Ronda 1 con 5 enemigos. Se muestra un anuncio 'RONDA 1'.",
    "Los enemigos aparecen uno a uno en los puntos de spawn y persiguen al jugador.",
    "El jugador dispara a los enemigos con clic izquierdo (gasta munici\u00f3n del cargador).",
    "Al quedarse sin balas, presiona R para recargar (2 segundos).",
    "Al eliminar todos los enemigos de la ronda, comienza la siguiente con m\u00e1s enemigos.",
    "La barra de vida del jugador se reduce al recibir da\u00f1o y se recupera con monedas de vida."
], 1):
    doc.add_paragraph(f"{i}. {s}", style='List Number')

h2(doc, "Caso de Uso: Sistema de Rondas")
para(doc, "Actor: Sistema")
para(doc, "Descripci\u00f3n: El sistema gestiona la progresi\u00f3n de rondas, incrementando la dificultad.")
para(doc, "Precondici\u00f3n: La ronda anterior ha sido completada (todos los enemigos eliminados).")
para(doc, "Secuencia normal:")
for i, s in enumerate([
    "El sistema detecta que no hay enemigos vivos ni pendientes de spawnear.",
    "Inicia la siguiente ronda (Ronda N+1).",
    "Calcula el total de enemigos: base + incremento * (ronda - 1).",
    "Muestra el anuncio 'RONDA N' en pantalla durante 3 segundos.",
    "Spawnnea los enemigos uno a uno con intervalo configurable.",
    "Espera a que todos los enemigos sean eliminados para repetir el ciclo."
], 1):
    doc.add_paragraph(f"{i}. {s}", style='List Number')

h2(doc, "Caso de Uso: Pausa")
para(doc, "Actor: Jugador")
para(doc, "Descripci\u00f3n: El jugador pausa el juego y elige entre reanudar, reiniciar o salir.")
para(doc, "Precondici\u00f3n: La partida est\u00e1 en curso.")
para(doc, "Secuencia normal:")
for i, s in enumerate([
    "El jugador presiona Escape.",
    "El juego se pausa y se muestra el men\u00fa de pausa.",
    "El jugador selecciona Reanudar para continuar.",
    "O selecciona Reiniciar para comenzar de nuevo.",
    "O selecciona Salir para volver al escritorio."
], 1):
    doc.add_paragraph(f"{i}. {s}", style='List Number')

h1(doc, "Diagrama de Clases")
para(doc, "Las clases principales del sistema y sus relaciones:")
classes_info = [
    "FPSController -> CharacterController (movimiento y mirada)",
    "Arma -> Camera (raycast de disparo y detecci\u00f3n de enemigo)",
    "Enemigo -> NavMeshAgent (persecuci\u00f3n), -> JugadorSalud (ataque)",
    "JugadorSalud <- Enemigo (recibe da\u00f1o), <- MonedaVida (recibe curaci\u00f3n)",
    "SpawnerEnemigos -> Enemigo (instanciaci\u00f3n), escucha evento OnMuerte",
    "Enemigo dispara evento OnMuerte -> SpawnerEnemigos (actualiza contador)",
    "GestorPausa controla Time.timeScale y Cursor.lockState"
]
for c in classes_info:
    doc.add_paragraph(c, style='List Bullet')

h1(doc, "Flujo de Escenas")
para(doc, "MenuInicio.unity -> MenuPrincipal.IniciarJuego() -> SceneManager.LoadScene('SampleScene')")
para(doc, "SampleScene.unity -> GestorPausa.Reiniciar() -> SceneManager.LoadScene(escena actual)")
para(doc, "SampleScene.unity -> JugadorSalud.ReiniciarJuego() -> SceneManager.LoadScene(escena actual)")

h1(doc, "Pruebas")
para(doc, "Las pruebas se ejecutan mediante Unity Test Framework (com.unity.test-framework 1.6.0) a trav\u00e9s del Test Runner del Editor de Unity. No existen scripts de prueba para l\u00ednea de comandos en el repositorio.")

h1(doc, "Referencias")
refs = [
    "Unity Technologies. Unity 6000.3.7f1 Documentation. https://docs.unity3d.com",
    "Microsoft. Introducci\u00f3n al lenguaje C# y .NET Framework. https://msdn.microsoft.com",
    "Sangucho Cueva, J. (2015). Metodolog\u00eda ICONIX para el desarrollo \u00e1gil de software."
]
for r in refs:
    doc.add_paragraph(r, style='List Bullet')

save(doc, "Tecnico.docx")

# ================================================================
# 3. USUARIO
# ================================================================
doc = Document()
add_title_page(doc, "Shooter1persona", "Manual de Usuario")

h1(doc, "Contenido")
idx = [
    "Introducci\u00f3n",
    "Objetivo del Juego",
    "Requisitos del Sistema",
    "Controles",
    "Men\u00fa Principal",
    "Interfaz de Juego",
    "Elementos del Juego",
    "Men\u00fa de Pausa",
    "Pantalla de Muerte"
]
for item in idx:
    doc.add_paragraph(item, style='List Bullet')

h1(doc, "Introducci\u00f3n")
para(doc, "Shooter1persona es un videojuego de disparos en primera persona (FPS) desarrollado en Unity como proyecto de ingenier\u00eda de sistemas. El jugador se enfrenta a oleadas de enemigos en un entorno 3D, poniendo a prueba su punter\u00eda y capacidad de supervivencia.")
para(doc, "El juego ofrece una experiencia din\u00e1mica donde los enemigos pueden aparecer en diferentes puntos del escenario y persiguen al jugador de manera inteligente usando navegaci\u00f3n por mallas (NavMesh). Algunos enemigos mutan a variantes m\u00e1s peque\u00f1as y r\u00e1pidas, aumentando la dificultad progresivamente.")

h1(doc, "Objetivo del Juego")
para(doc, "Sobrevivir el mayor tiempo posible eliminando a los enemigos que aparecen en oleadas. El contador de bajas (KILLS) registra el progreso del jugador. La partida termina cuando la salud del jugador llega a cero.")

h1(doc, "Requisitos del Sistema")
for req in [
    "Sistema Operativo: Windows 10 u 11 (64-bit)",
    "Procesador: Intel Core i3-2100 o superior",
    "Memoria RAM: 8 GB",
    "Tarjeta Gr\u00e1fica: Compatible con DirectX 11 o superior",
    "Almacenamiento: 1 GB de espacio disponible"
]:
    doc.add_paragraph(req, style='List Bullet')

h1(doc, "Controles")
controls_table = doc.add_table(rows=8, cols=2)
controls_table.style = 'Light Grid Accent 1'
controls_table.alignment = WD_TABLE_ALIGNMENT.CENTER
controls_table.rows[0].cells[0].text = 'Tecla'
controls_table.rows[0].cells[1].text = 'Acci\u00f3n'
controls_data = [
    ('W, A, S, D', 'Moverse adelante, izquierda, atr\u00e1s, derecha'),
    ('Shift (izquierdo)', 'Correr'),
    ('Rat\u00f3n (movimiento)', 'Mirar alrededor'),
    ('Clic izquierdo', 'Disparar'),
    ('R', 'Recargar arma'),
    ('Escape', 'Pausar / Reanudar'),
    ('Apuntar a un enemigo', 'Muestra la barra de vida del enemigo'),
]
for i, (key, action) in enumerate(controls_data):
    controls_table.rows[i+1].cells[0].text = key
    controls_table.rows[i+1].cells[1].text = action

doc.add_paragraph()

h1(doc, "Men\u00fa Principal")
para(doc, "Al iniciar el juego aparece el men\u00fa principal con las siguientes opciones:")
doc.add_paragraph("Jugar: Inicia la partida en la escena de juego.", style='List Bullet')
doc.add_paragraph("Salir: Cierra la aplicaci\u00f3n.", style='List Bullet')

h1(doc, "Interfaz de Juego")
para(doc, "Durante la partida, la pantalla muestra los siguientes elementos:")
doc.add_paragraph("Barra de vida (esquina superior izquierda): Muestra la salud actual del jugador mediante un Slider. Se reduce al recibir da\u00f1o y se recupera con monedas de vida.", style='List Bullet')
doc.add_paragraph("Municion (junto a la barra de vida): Muestra las balas en el cargador y la capacidad total (ej. '27 / 30').", style='List Bullet')
doc.add_paragraph("Contador KILLS (esquina superior derecha): Muestra el n\u00famero de enemigos eliminados en toda la partida.", style='List Bullet')
doc.add_paragraph("Anuncio de Ronda (centro): Aparece 'RONDA 1', 'RONDA 2', etc. al inicio de cada ronda.", style='List Bullet')
doc.add_paragraph("Mira (centro): Indica hacia d\u00f3nde apunta el arma.", style='List Bullet')
doc.add_paragraph("Barra de vida del enemigo: Aparece sobre el enemigo cuando el jugador lo mira directamente.", style='List Bullet')
doc.add_paragraph("Minimapa: Muestra una vista cenital de la posici\u00f3n del jugador en el escenario.", style='List Bullet')

h1(doc, "Sistema de Rondas")
para(doc, "El juego se organiza en rondas progresivas. La Ronda 1 comienza con 5 enemigos y cada ronda siguiente suma 3 enemigos adicionales. Al inicio de cada ronda aparece un anuncio en pantalla. La ronda termina cuando todos los enemigos han sido eliminados, y tras una breve pausa comienza la siguiente.")

h1(doc, "Elementos del Juego")
h2(doc, "Enemigos")
para(doc, "Los enemigos aparecen en puntos de spawn distribuidos por el escenario y persiguen al jugador autom\u00e1ticamente. Tienen un 30% de probabilidad de ser variantes peque\u00f1as y r\u00e1pidas con menor salud pero mayor velocidad y da\u00f1o.")

h2(doc, "Municion")
para(doc, "El arma tiene un cargador de 30 balas y una reserva total de 120 balas. Al disparar se consume una bala del cargador. Cuando el cargador se vac\u00eda, se puede recargar presionando la tecla R. La recarga tarda 2 segundos y consume balas de la reserva. Si se dispara con el cargador vac\u00edo y hay balas en reserva, la recarga se inicia autom\u00e1ticamente.")

h2(doc, "Monedas de Vida")
para(doc, "Al eliminar enemigos, existe una probabilidad del 20% de que suelten una moneda de vida. Al pasar sobre ella, el jugador recupera 20 puntos de salud.")

h1(doc, "Men\u00fa de Pausa")
para(doc, "Presionando la tecla Escape durante la partida se abre el men\u00fa de pausa (el juego se congela y el cursor se desbloquea). Opciones:")
doc.add_paragraph("Reanudar: Contin\u00faa la partida desde donde se qued\u00f3 (restaura el tiempo y bloquea el cursor).", style='List Bullet')
doc.add_paragraph("Reiniciar: Reinicia la partida desde el principio, recargando toda la escena.", style='List Bullet')
doc.add_paragraph("Salir: Cierra la aplicaci\u00f3n por completo.", style='List Bullet')

h1(doc, "Pantalla de Muerte")
para(doc, "Cuando la salud del jugador llega a cero, se muestra una pantalla de muerte con la opci\u00f3n de reiniciar la partida haciendo clic. El cursor se desbloquea para permitir la interacci\u00f3n.")

save(doc, "Usuario.docx")

print("All 3 documents created successfully!")
